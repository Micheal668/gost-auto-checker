from __future__ import annotations

from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any

from docx import Document


def write_result(media_root: Path | str, job_id: str, findings: List[Dict[str, Any]],snapshot: Dict[str, Any] | None = None) -> str:
    """
    生成结果 docx，写入 MEDIA_ROOT/results/ 下
    返回：相对 MEDIA_ROOT 的路径，例如：results/gost_result_<job_id>.docx
    """
    media_root = Path(media_root)
    out_dir = media_root / "results"
    out_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Результаты проверки отчёта по ГОСТ 7.32-2017 (MVP)", level=1)
    doc.add_paragraph(f"Job ID: {job_id}")
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    # ============================
    # ✅ 结构锚点清单（即使 issues=0 也会显示）
    # ============================
    # if snapshot and snapshot.get("anchor_map"):
    #     doc.add_heading("Структурные элементы (якоря)", level=2)

    # # 按段落 idx 排序，输出稳定
    #     items = sorted(snapshot["anchor_map"].items(), key=lambda kv: kv[1])
    #     for title, idx in items:
    #         doc.add_paragraph(f"{title} #{idx}")


    table = doc.add_table(rows=1, cols=4)
    
    hdr = table.rows[0].cells
    # hdr[0].text = "Стр./позиция"
    # hdr[1].text = "Уровень"
    # hdr[2].text = "Ошибка"
    # hdr[3].text = "Рекомендация (AI/шаблон)"
    hdr[0].text = "Уровень"
    hdr[1].text = "Ошибка"
    hdr[2].text = "Рекомендация (AI/шаблон)"
    
    for fin in (findings or []):

        # page
        anchor = (fin.get("anchor") or "").strip()
        para_idx = fin.get("para_idx")
        if anchor or para_idx is not None:
            page = f"{anchor} #{para_idx if para_idx is not None else '?'}".strip()
        else:
            page = str(fin.get("page", "?"))
        # sev
        sev = str(fin.get("severity", "NEED_REVIEW"))

        # 把 rule_id/clause/category 内嵌到“错误”列，保证可追溯
        rid = fin.get("rule_id")
        clause = fin.get("clause")
        cat = fin.get("category")

        # err-text
        prefix_parts = []
        if rid:
            prefix_parts.append(f"[{rid}]")
        if clause:
            prefix_parts.append(f"§{clause}")
        if cat:
            prefix_parts.append(str(cat))

        prefix = " ".join(prefix_parts)
        msg = str(fin.get("message", ""))
        err_text = f"{prefix} {msg}".strip()

        
        suggestion = str(fin.get("suggestion", ""))

        row = table.add_row().cells
        # row[0].text = page
        # row[1].text = sev
        # row[2].text = err_text
        # row[3].text = suggestion
        row[0].text = sev
        row[1].text = err_text
        row[2].text = suggestion


    filename = f"gost_result_{job_id}.docx"
    out_path = out_dir / filename
    doc.save(out_path)

    # ✅ 返回相对 MEDIA_ROOT 的路径：results/xxx.docx
    return str(Path("results") / filename)
